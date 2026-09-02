"""RAG 核心引擎 - 文档加载、分块、向量化、检索

设计要点:
- chunk ID 基于内容哈希 + upsert，重复/更新的文档不会静默丢失
- 页码元数据贯穿始终（pymupdf4llm page_chunks，1-based；非 PDF 格式无页码）
- 检索返回相似度分数并按阈值过滤，全被过滤时附带提示与最高分供 Agent 决策
- 文档路径统一归一化（normcase+normpath），避免 Windows 大小写/分隔符差异
  导致同一文档重复入库或删除失效
- 日志只写 stderr，绝不碰 stdout（stdio MCP 协议通道）
"""

import hashlib
import json
import logging
import os
import threading
import time
from typing import Any

import chromadb
import pymupdf4llm
from chromadb.config import Settings
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config import Config

logger = logging.getLogger("rag_engine")

# 中文感知的分块分隔符，按优先级递减
_SEPARATORS = ["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""]

# 支持的文档格式
SUPPORTED_EXTENSIONS = frozenset({".pdf", ".txt", ".md", ".markdown", ".docx"})


class RAGEngine:
    """RAG 引擎：文档加载、分块、向量化与检索（不含 LLM 生成）"""

    COLLECTION_NAME = "documents"

    def __init__(self):
        """初始化 ChromaDB 客户端、分块器与写入锁"""
        self.chroma_client = chromadb.PersistentClient(
            path=Config.CHROMA_PERSIST_DIR, settings=Settings(anonymized_telemetry=False)
        )
        self.collection = self.chroma_client.get_or_create_collection(
            name=self.COLLECTION_NAME,
            metadata={"hnsw:space": "cosine", "embedding_model": Config.EMBEDDING_MODEL},
        )

        # 若知识库是用其他 embedding 模型构建的，检索维度会不匹配，尽早报错
        stored_model = (self.collection.metadata or {}).get("embedding_model")
        if stored_model and stored_model != Config.EMBEDDING_MODEL:
            raise RuntimeError(
                f"知识库由 embedding 模型 {stored_model!r} 构建，当前配置为 "
                f"{Config.EMBEDDING_MODEL!r}，向量维度可能不匹配。"
                f"请清空知识库（clear_knowledge_base）或改回原模型。"
            )

        self._splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            separators=_SEPARATORS,
        )
        # Chroma 写操作互斥（工具经 asyncio.to_thread 在线程池并发执行）
        self._write_lock = threading.Lock()

        logger.info("RAGEngine 初始化完成，当前文档块数量: %d", self.collection.count())

    # ------------------------------------------------------------------ #
    # 文档加载
    # ------------------------------------------------------------------ #

    def load_document(self, file_path: str) -> str:
        """加载文档（PDF/TXT/MD/DOCX）：解析 -> 分块 -> 向量化 -> upsert 入库

        同一路径同一内容: 跳过（提示未变化）；同一路径内容有更新: 先删旧块再入库。
        """
        if not file_path or not file_path.strip():
            return "错误: file_path 不能为空"
        path = _normalize_source(file_path)
        if not _has_supported_extension(path):
            supported = "、".join(sorted(SUPPORTED_EXTENSIONS))
            return f"错误: 不支持的文件格式 - {path}（支持: {supported}）"

        file_hash = _sha256_file(path)
        logger.info("开始加载文档: %s (sha256=%s)", path, file_hash[:12])

        with self._write_lock:
            existing = self._chunks_of_source(path)
            if existing:
                unchanged = all(m.get("doc_id") == file_hash for m, _ in existing)
                if unchanged:
                    logger.info("文档未变化，跳过: %s", path)
                    return (
                        f"文档已存在于知识库且内容未变化，共 {len(existing)} 个文档块，"
                        f"本次跳过: {_display_name(path)}"
                    )
                # 内容有更新: 删除旧版本全部块（含历史路径写法不同的残留版本）
                self._delete_by_source(path)
                logger.info("检测到文档更新，已删除旧版本 %d 个块", len(existing))

            n_chunks = self._ingest(path, file_hash)

        return (
            f"成功加载文档: {_display_name(path)}，共 {n_chunks} 个文档块，"
            f"知识库当前总计 {self.collection.count()} 个块"
        )

    def _ingest(self, path: str, file_hash: str) -> int:
        """解析文档、逐页分块、批量向量化并 upsert。调用方需持有 _write_lock。"""
        ids: list[str] = []
        embeddings: list[list[float]] = []
        metadatas: list[dict[str, Any]] = []
        documents: list[str] = []

        for page_no, page_text in _extract_pages(path):
            if not page_text:
                continue
            for idx, chunk in enumerate(self._splitter.split_text(page_text)):
                chunk = chunk.strip()
                if not chunk:
                    continue
                ids.append(_chunk_id(path, file_hash, chunk))
                metadata = {"source": path, "doc_id": file_hash, "chunk_index": idx}
                if page_no is not None:  # ChromaDB 元数据不支持 null，无页码格式直接省略
                    metadata["page"] = page_no
                metadatas.append(metadata)
                documents.append(chunk)

        if not documents:
            raise ValueError(
                f"无法从文档提取到文本: {_display_name(path)}（PDF 可能是扫描件，或文档为空）"
            )

        # 分批向量化（text-embedding-v4 单次最多 10 条）
        for i in range(0, len(documents), Config.EMBED_BATCH_SIZE):
            batch = documents[i : i + Config.EMBED_BATCH_SIZE]
            embeddings.extend(_embed_batch(batch))

        self.collection.upsert(ids=ids, embeddings=embeddings, metadatas=metadatas, documents=documents)
        logger.info("入库完成: %s, %d 个块", _display_name(path), len(documents))
        return len(documents)

    # ------------------------------------------------------------------ #
    # 检索
    # ------------------------------------------------------------------ #

    def search(self, question: str, top_k: int | None = None) -> dict[str, Any]:
        """向量检索知识库，返回带相似度分数的文档块（低于阈值的结果被过滤）"""
        question = (question or "").strip()
        if not question:
            return {"results": [], "count": 0, "error": "问题不能为空"}
        if self.collection.count() == 0:
            return {"results": [], "count": 0, "error": "知识库为空，请先使用 load_document 加载文档"}

        k = max(1, min(top_k or Config.TOP_K, 20))
        query_embedding = _embed_batch([question])[0]

        raw = self.collection.query(
            query_embeddings=[query_embedding], n_results=min(k, self.collection.count())
        )

        results: list[dict[str, Any]] = []
        similarities: list[float] = []
        metadatas = (raw.get("metadatas") or [[]])[0]
        documents = (raw.get("documents") or [[]])[0]
        distances = (raw.get("distances") or [[]])[0]

        for metadata, text, distance in zip(metadatas, documents, distances):
            similarity = round(1.0 - float(distance), 4)  # cosine distance -> similarity
            similarities.append(similarity)
            if similarity < Config.SCORE_THRESHOLD:
                continue
            results.append(
                {
                    "text": text,
                    "source": metadata.get("source", ""),
                    "page": metadata.get("page"),
                    "similarity": similarity,
                    "chunk_index": metadata.get("chunk_index"),
                }
            )

        results.sort(key=lambda r: r["similarity"], reverse=True)
        payload: dict[str, Any] = {"results": results, "count": len(results), "error": None}
        if not results and similarities:
            # 有候选但全被阈值过滤：附上最高分，帮调用方区分"不相关"与"库为空"
            best = max(similarities)
            payload["best_similarity"] = best
            payload["hint"] = (
                f"检索到 {len(similarities)} 个候选块，但最高相似度 {best} 仍低于阈值 "
                f"{Config.SCORE_THRESHOLD}，已全部过滤。知识库中可能没有与该问题相关的内容。"
            )
        return payload

    # ------------------------------------------------------------------ #
    # 文档管理
    # ------------------------------------------------------------------ #

    def list_documents(self) -> dict[str, Any]:
        """列出知识库中的所有文档及其块数、页码范围"""
        if self.collection.count() == 0:
            return {"documents": [], "count": 0}

        all_meta = self.collection.get(include=["metadatas"]).get("metadatas") or []
        docs: dict[str, dict[str, Any]] = {}
        for m in all_meta:
            source = m.get("source", "<unknown>")
            entry = docs.setdefault(
                source,
                {"source": source, "doc_id": m.get("doc_id", ""), "chunks": 0, "pages": set()},
            )
            entry["chunks"] += 1
            if m.get("page") is not None:
                entry["pages"].add(m["page"])

        documents = [
            {**d, "pages": sorted(d["pages"])} for d in sorted(docs.values(), key=lambda x: x["source"])
        ]
        return {"documents": documents, "count": len(documents)}

    def delete_document(self, source: str) -> str:
        """按源文件路径删除单个文档的所有块（路径匹配忽略大小写与分隔符差异）"""
        source = (source or "").strip()
        if not source:
            return "错误: source 不能为空（可先用 list_documents 查看准确的文件路径）"

        with self._write_lock:
            existing = self._chunks_of_source(source)
            if not existing:
                return f"知识库中不存在该文档: {source}"
            self._delete_by_source(source)

        return f"已删除文档 {_display_name(source)}，共 {len(existing)} 个文档块"

    def clear(self) -> str:
        """清空知识库（删除并重建集合）"""
        with self._write_lock:
            try:
                self.chroma_client.delete_collection(self.COLLECTION_NAME)
            except Exception:  # 集合不存在时忽略
                logger.debug("集合不存在，无需删除", exc_info=True)
            self.collection = self.chroma_client.get_or_create_collection(
                name=self.COLLECTION_NAME,
                metadata={"hnsw:space": "cosine", "embedding_model": Config.EMBEDDING_MODEL},
            )
        logger.info("知识库已清空")
        return "知识库已清空"

    # ------------------------------------------------------------------ #
    # 内部工具
    # ------------------------------------------------------------------ #

    def _chunks_of_source(self, source: str) -> list[tuple[dict[str, Any], str]]:
        """取某一路径下现存的所有块（归一化匹配），返回 [(metadata, id), ...]

        全量扫描而非 where 精确过滤：ChromaDB 元数据过滤只支持字符串精确匹配，
        无法兼容历史数据中大小写/分隔符不同的同一路径。个人知识库规模下可接受。
        """
        target = _normalize_source(source)
        got = self.collection.get(include=["metadatas"])
        return [
            (m, i)
            for m, i in zip(got.get("metadatas") or [], got.get("ids") or [])
            if _normalize_source(m.get("source", "")) == target
        ]

    def _delete_by_source(self, source: str) -> None:
        """删除某一路径下所有块。ChromaDB 按存储值过滤，需逐个历史写法删除。"""
        for stored in {m.get("source", "") for m, _ in self._chunks_of_source(source)}:
            self.collection.delete(where={"source": stored})


# ---------------------------------------------------------------------- #
# 模块级工具函数（不依赖实例状态，便于单测）
# ---------------------------------------------------------------------- #


def _normalize_source(path: str) -> str:
    """归一化文档路径：统一分隔符、折叠 .与..，Windows 下额外忽略大小写

    ChromaDB 元数据过滤是精确字符串匹配，不归一化的话同一文件换个写法
    （D:/a.pdf vs d:\a.PDF）会被当成两个文档，产生重复入库或删除失效。
    """
    return os.path.normcase(os.path.normpath(path.strip()))


def _has_supported_extension(path: str) -> bool:
    """按扩展名粗判格式（详细校验交给各解析器）"""
    return os.path.splitext(path)[1].lower() in SUPPORTED_EXTENSIONS


def _sha256_file(path: str) -> str:
    """计算文件内容 sha256，用于文档去重与更新检测"""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(1 << 20), b""):
            h.update(block)
    return h.hexdigest()


def _chunk_id(source: str, file_hash: str, chunk: str) -> str:
    """内容寻址的 chunk ID：同内容同 ID（配合 upsert 幂等），不同路径互不干扰"""
    digest = hashlib.sha256(f"{source}|{file_hash}|{chunk}".encode("utf-8")).hexdigest()
    return f"{file_hash[:12]}_{digest[:32]}"


def _display_name(path: str) -> str:
    return os.path.basename(path)


def _extract_pages(path: str) -> list[tuple[int | None, str]]:
    """按扩展名解析文档，返回 [(page_no, text), ...]

    page_no 为 1-based 页码，仅 PDF 有分页概念；其余格式返回 None
    （ChromaDB 元数据不支持 null，入库时会省略该键）。
    """
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return [(None, _extract_docx(path))]
    return [(None, _read_text(path))]


def _extract_pdf(path: str) -> list[tuple[int | None, str]]:
    """PDF -> Markdown，逐页返回（pymupdf4llm 的 page_number 已是 1 基）"""
    pages = pymupdf4llm.to_markdown(path, page_chunks=True)
    return [
        (int(item["metadata"].get("page_number", 1)), (item.get("text") or "").strip())
        for item in pages
    ]


def _extract_docx(path: str) -> str:
    """DOCX -> 纯文本：正文段落按行拼接，表格逐行用 | 连接（python-docx 延迟导入）"""
    from docx import Document

    doc = Document(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_text(path: str) -> str:
    """读取纯文本文件，编码依次尝试 utf-8-sig(BOM)/utf-8/gbk（兼容中文 Windows 文件）"""
    with open(path, "rb") as f:
        raw = f.read()
    for encoding in ("utf-8-sig", "utf-8", "gbk"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError(f"无法识别文件编码（已尝试 utf-8/gbk）: {_display_name(path)}")


def _embed_batch(texts: list[str]) -> list[list[float]]:
    """调用 DashScope embedding，带指数退避重试"""
    from dashscope import TextEmbedding

    last_error: Exception | None = None
    for attempt in range(Config.EMBED_MAX_RETRIES):
        try:
            response = TextEmbedding.call(model=Config.EMBEDDING_MODEL, input=texts)
            if response.status_code != 200:
                raise RuntimeError(f"Embedding API 返回 {response.status_code}: {response.message}")
            return [item["embedding"] for item in response.output["embeddings"]]
        except Exception as e:  # noqa: BLE001 - 统一记日志后重试
            last_error = e
            wait = 2**attempt
            logger.warning("Embedding 调用失败（第 %d 次）: %s，%ds 后重试", attempt + 1, e, wait)
            time.sleep(wait)
    raise RuntimeError(f"Embedding API 连续 {Config.EMBED_MAX_RETRIES} 次调用失败: {last_error}")


def results_to_text(payload: dict[str, Any]) -> str:
    """把检索结果序列化为给 MCP 客户端阅读的 JSON 文本"""
    return json.dumps(payload, ensure_ascii=False, indent=2)
